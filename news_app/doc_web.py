import json
import os
from collections import Counter
from docx import Document
from docx.shared import Pt
from io import BytesIO
from flask import jsonify

def load_json_files(folder_path):
    data = []
    for filename in os.listdir(folder_path):
        if filename.endswith(".json"):
            with open(os.path.join(folder_path, filename), "r", encoding="utf-8") as f:
                data.append(json.load(f))
    return data

def analyze_emotions(data):
    emotion_counter = Counter()
    for item in data:
        for quad in item.get("四元组", []):
            emotion = quad.get("情感", "中性")
            emotion_counter[emotion] += 1
    total = sum(emotion_counter.values())
    distribution = {k: f"{v}条，占比{v/total:.0%}" for k, v in emotion_counter.items()}
    return emotion_counter, distribution, total

def extract_example(data, target_emotion):
    for item in data:
        for quad in item.get("四元组", []):
            if quad.get("情感") == target_emotion:
                return item["title"], item["link"]
    return "无代表性舆情", "#"

#本地版，已经弃用
def generate_report(news_keyword, output_path="1汕头城市舆情分析报告.docx"):
    json_folder = os.path.join("./news_app/shantou_city", news_keyword)
    data = load_json_files(json_folder)
    emotion_counts, emotion_distribution, total_count = analyze_emotions(data)

    positive_title, positive_link = extract_example(data, "正面")
    negative_title, negative_link = extract_example(data, "负面")
    neutral_title, neutral_link = extract_example(data, "中性")

    doc = Document()
    doc.add_heading("汕头城市舆情分析报告", 0)

    # 一、数据概况
    doc.add_heading("一、数据概况", level=1)
    doc.add_paragraph(
        f"本次舆情分析共收集到{total_count}条与汕头相关的舆情源数据。"
        f"从情感类型分布来看，" +
        "，".join([f"{k}舆情{v}" for k, v in emotion_distribution.items()]) +
        "。数据均来自联网搜索平台。"
    )

    # 二、整体情感倾向分析
    doc.add_heading("二、整体情感倾向分析", level=1)
    doc.add_heading("核心结论总结", level=2)
    doc.add_paragraph("汕头城市舆情整体呈现积极态势，正面舆情占据主导地位。")
    doc.add_heading("论据呈现及结论推导", level=2)
    doc.add_paragraph(f"正面舆情示例：[{positive_title}]({positive_link})")
    doc.add_paragraph(f"中性舆情示例：[{neutral_title}]({neutral_link})")
    doc.add_paragraph(f"负面舆情示例：[{negative_title}]({negative_link})")
    doc.add_heading("策略建议", level=2)
    doc.add_paragraph("- 加大正面舆情的传播；\n- 回应负面舆论，澄清事实；\n- 挖掘中性信息中的正向价值。")

    # 三、负面舆情分析
    doc.add_heading("三、负面舆情分析", level=1)
    doc.add_heading("核心结论总结", level=2)
    doc.add_paragraph("汕头负面舆情主要涉及城市发展问题，虽然占比小，但仍需关注。")
    doc.add_heading("论据呈现及结论推导", level=2)
    doc.add_paragraph(f"例如：[负面舆情示例]({negative_link}) 反映了公众关注的某些问题。")
    doc.add_heading("策略建议", level=2)
    doc.add_paragraph("- 加强城市管理，提高服务质量。\n- 明确宣传城市发展方向，化解误解。")

    # 四、正面舆情传播效果分析（简化）
    doc.add_heading("四、正面舆情传播效果分析", level=1)
    doc.add_heading("核心结论总结", level=2)
    doc.add_paragraph("正面舆情传播广度一般，用户互动仍有提升空间。")
    doc.add_heading("策略建议", level=2)
    doc.add_paragraph("- 增强内容吸引力；\n- 推动内容在社交平台二次传播。")

    # 五、总结与展望
    doc.add_heading("五、总结与展望", level=1)
    doc.add_paragraph(
        "总体而言，汕头城市舆情整体积极，但仍存在一些负面因素和传播效果不佳的问题。"
        "未来应加强城市建设与舆情管理，提升形象，吸引资源与关注。"
    )

    doc.save(output_path)
    print(f"报告已生成：{output_path}")

#联网版
def generate_report_to_bytes(news_keyword):
    # 模拟读取并分析数据
    json_folder = os.path.normpath(os.path.join("shantou_city", news_keyword))
    
    data = load_json_files(json_folder)
    emotion_counts, emotion_distribution, total_count = analyze_emotions(data)
    positive_title, positive_link = extract_example(data, "正面")
    negative_title, negative_link = extract_example(data, "负面")
    neutral_title, neutral_link = extract_example(data, "中性")

    doc = Document()
    doc.add_heading("汕头城市舆情分析报告", 0)
    doc.add_heading("一、数据概况", level=1)
    doc.add_paragraph(
        f"本次舆情分析共收集到{total_count}条与汕头相关的舆情源数据。"
        f"从情感类型分布来看，" +
        "，".join([f"{k}舆情{v}" for k, v in emotion_distribution.items()]) +
        "。数据均来自联网搜索平台。"
    )
    doc.add_heading("二、整体情感倾向分析", level=1)
    doc.add_heading("核心结论总结", level=2)
    doc.add_paragraph("汕头城市舆情整体呈现积极态势，正面舆情占据主导地位。")
    doc.add_heading("论据呈现及结论推导", level=2)
    doc.add_paragraph(f"正面舆情示例：[{positive_title}]({positive_link})")
    doc.add_paragraph(f"中性舆情示例：[{neutral_title}]({neutral_link})")
    doc.add_paragraph(f"负面舆情示例：[{negative_title}]({negative_link})")
    doc.add_heading("策略建议", level=2)
    doc.add_paragraph("- 加大正面舆情的传播；\n- 回应负面舆论，澄清事实；\n- 挖掘中性信息中的正向价值。")

    doc.add_heading("三、负面舆情分析", level=1)
    doc.add_heading("核心结论总结", level=2)
    doc.add_paragraph("汕头负面舆情主要涉及城市发展问题，虽然占比小，但仍需关注。")
    doc.add_heading("论据呈现及结论推导", level=2)
    doc.add_paragraph(f"例如：[负面舆情示例]({negative_link}) 反映了公众关注的某些问题。")
    doc.add_heading("策略建议", level=2)
    doc.add_paragraph("- 加强城市管理，提高服务质量。\n- 明确宣传城市发展方向，化解误解。")

    doc.add_heading("四、正面舆情传播效果分析", level=1)
    doc.add_heading("核心结论总结", level=2)
    doc.add_paragraph("正面舆情传播广度一般，用户互动仍有提升空间。")
    doc.add_heading("策略建议", level=2)
    doc.add_paragraph("- 增强内容吸引力；\n- 推动内容在社交平台二次传播。")

    doc.add_heading("五、总结与展望", level=1)
    doc.add_paragraph(
        "总体而言，汕头城市舆情整体积极，但仍存在一些负面因素和传播效果不佳的问题。"
        "未来应加强城市建设与舆情管理，提升形象，吸引资源与关注。"
    )

    # 将 Word 写入内存中
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io

# 示例调用
# generate_report("your_json_folder_path")
if __name__ == "__main__":
    # 替换为实际的JSON文件夹路径
    json_folder = "潮南区"
    generate_report(json_folder)